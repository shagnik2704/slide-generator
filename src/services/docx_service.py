"""
Word Document export/import service for Spoken Tutorial scripts.
Creates editable .docx files with two-column table format (Visual Cue | Narration).
"""
import re
from pathlib import Path
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def json_to_docx(json_data: dict, output_path: str = None) -> BytesIO:
    """
    Convert JSON script to a Word document with two-column table.
    
    Args:
        json_data: The JSON script data with slides, metadata, etc.
        output_path: Optional path to save the file. If None, returns BytesIO buffer.
    
    Returns:
        BytesIO buffer containing the .docx file, or path if output_path provided.
    """
    doc = Document()
    
    # Add title
    title = json_data.get('presentation_title', 'Spoken Tutorial Script')
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata section
    _add_metadata_section(doc, json_data)
    
    # Add script table header
    doc.add_heading('Script', level=1)
    doc.add_paragraph('Edit the narration in the right column. Do not modify the table structure.')
    
    # Create two-column table
    slides = json_data.get('slides', [])
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.5)
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Visual Cue'
    header_cells[1].text = 'Narration'
    
    # Style header row
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)
    
    # Add slide rows
    for i, slide in enumerate(slides):
        row = table.add_row()
        
        # Visual Cue column (slide title + image prompt)
        visual_cue = _format_visual_cue(slide)
        row.cells[0].text = visual_cue
        
        # Narration column (with bold markdown parsing)
        narration = slide.get('narration', '')
        # Convert \n to actual newlines for display
        narration = narration.replace('\\n', '\n')
        _add_formatted_text(row.cells[1], narration)
        
        # Add padding to all cells in this row
        for cell in row.cells:
            _set_cell_padding(cell, top=100, bottom=100, left=100, right=100)
    
    # Save or return buffer
    if output_path:
        doc.save(output_path)
        return output_path
    else:
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


def _set_cell_padding(cell, top=0, bottom=0, left=0, right=0):
    """Set cell padding/margins in twips (1/20 of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for side, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    
    tcPr.append(tcMar)


def _add_formatted_text(cell, text):
    """Add text to cell with **bold** markdown parsed as actual bold."""
    from docx.shared import Pt as PtSpacing
    
    # Clear existing text
    cell.text = ""
    paragraph = cell.paragraphs[0]
    
    # Set line spacing (1.5 lines = 1.5 * 12pt = 18pt)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)  # Extra space after paragraph
    
    # Parse **bold** pattern
    pattern = r'\*\*(.+?)\*\*'
    last_end = 0
    
    for match in re.finditer(pattern, text):
        # Add normal text before this match
        if match.start() > last_end:
            normal_text = text[last_end:match.start()]
            paragraph.add_run(normal_text)
        
        # Add bold text
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        
        last_end = match.end()
    
    # Add remaining normal text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def docx_to_json(docx_file) -> dict:
    """
    Parse a Word document back to JSON script format.
    
    Args:
        docx_file: File path, file object, or BytesIO buffer of the .docx file.
    
    Returns:
        Dictionary with the parsed script data.
    """
    doc = Document(docx_file)
    
    # Extract title from first heading
    title = ""
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            title = para.text
            break
    
    # Find the main script table (should be the largest table)
    script_table = None
    for table in doc.tables:
        # Check if it looks like our script table (has Visual Cue and Narration headers)
        if len(table.rows) > 0 and len(table.columns) >= 2:
            first_row = table.rows[0]
            if 'Visual Cue' in first_row.cells[0].text or 'Narration' in first_row.cells[1].text:
                script_table = table
                break
    
    if not script_table:
        raise ValueError("Could not find script table in document. Make sure the table structure is preserved.")
    
    # Parse slides from table (skip header row)
    slides = []
    for row in script_table.rows[1:]:  # Skip header
        if len(row.cells) < 2:
            continue
            
        visual_cue = row.cells[0].text.strip()
        narration = row.cells[1].text.strip()
        
        # Skip empty rows
        if not visual_cue and not narration:
            continue
        
        # Extract title and image_prompt from visual cue
        slide_title, image_prompt = _parse_visual_cue(visual_cue)
        
        slides.append({
            'title': slide_title,
            'narration': narration,
            'image_prompt': image_prompt
        })
    
    # Try to extract metadata from document
    metadata = _extract_metadata(doc)
    
    return {
        'presentation_title': metadata.get('title', title),
        'module': metadata.get('module', ''),
        'episode': metadata.get('episode', ''),
        'duration': metadata.get('duration', '3-4 min'),
        'learning_objectives': metadata.get('learning_objectives', []),
        'prerequisites': metadata.get('prerequisites', ''),
        'meta_tags': metadata.get('meta_tags', []),
        'outline': metadata.get('outline', []),
        'slides': slides
    }


def _add_metadata_section(doc, json_data: dict):
    """Add metadata section to the document."""
    doc.add_heading('Metadata', level=1)
    
    # Create a simple metadata table
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.style = 'Table Grid'
    meta_table.columns[0].width = Inches(2)
    meta_table.columns[1].width = Inches(5)
    
    metadata_items = [
        ('Title', json_data.get('presentation_title', '')),
        ('Module', json_data.get('module', '')),
        ('Episode', json_data.get('episode', '')),
        ('Duration', json_data.get('duration', '')),
        ('Prerequisites', json_data.get('prerequisites', '')),
        ('Learning Objectives', '\n'.join(json_data.get('learning_objectives', []))),
    ]
    
    for label, value in metadata_items:
        if value:
            row = meta_table.add_row()
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = str(value)
    
    doc.add_paragraph()  # Add spacing


def _format_visual_cue(slide: dict) -> str:
    """Format the visual cue column for a slide."""
    title = slide.get('title', '')
    image_prompt = slide.get('image_prompt', '')
    
    # Combine title and image prompt
    parts = []
    if title:
        parts.append(f"[{title}]")
    if image_prompt and image_prompt != title:
        parts.append(image_prompt)
    
    return '\n'.join(parts) if parts else ''


def _parse_visual_cue(visual_cue: str) -> tuple:
    """Parse visual cue back to title and image_prompt."""
    lines = visual_cue.strip().split('\n')
    
    title = ''
    image_prompt = ''
    
    for line in lines:
        line = line.strip()
        # Check for [Title] format
        match = re.match(r'\[(.+)\]', line)
        if match:
            title = match.group(1)
        else:
            image_prompt = line
    
    # If no explicit title found, use the whole visual cue
    if not title and visual_cue:
        title = visual_cue.split('\n')[0].strip()
    
    # If no image_prompt, use title as image_prompt
    if not image_prompt:
        image_prompt = title
    
    return title, image_prompt


def _extract_metadata(doc) -> dict:
    """Extract metadata from document tables."""
    metadata = {}
    
    for table in doc.tables:
        # Look for metadata table (has Label | Value format)
        for row in table.rows:
            if len(row.cells) >= 2:
                label = row.cells[0].text.strip().lower()
                value = row.cells[1].text.strip()
                
                if 'title' in label and not metadata.get('title'):
                    metadata['title'] = value
                elif 'module' in label:
                    metadata['module'] = value
                elif 'episode' in label or 'tutorial' in label:
                    metadata['episode'] = value
                elif 'duration' in label:
                    metadata['duration'] = value
                elif 'prerequisite' in label:
                    metadata['prerequisites'] = value
                elif 'learning' in label and 'objective' in label:
                    # Split by newlines to get list
                    metadata['learning_objectives'] = [
                        obj.strip() for obj in value.split('\n') if obj.strip()
                    ]
    
    return metadata


def export_script_docx(json_data: dict, output_dir: str = None) -> dict:
    """
    Export JSON script to a downloadable .docx file.
    
    Args:
        json_data: The JSON script data
        output_dir: Directory to save the file. If None, uses static/ directory.
    
    Returns:
        Dictionary with 'file_path' and 'file_name'
    """
    import time
    
    # Generate filename
    project_root = Path(__file__).parent.parent.parent
    if output_dir is None:
        output_dir = project_root / "static"
    else:
        output_dir = Path(output_dir)
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate unique filename
    timestamp = int(time.time())
    title_slug = json_data.get('presentation_title', 'script').replace(' ', '_')[:30]
    filename = f"{title_slug}_{timestamp}.docx"
    file_path = output_dir / filename
    
    # Generate document
    json_to_docx(json_data, str(file_path))
    
    return {
        "file_path": str(file_path),
        "file_name": filename
    }
