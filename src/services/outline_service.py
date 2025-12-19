"""
Outline generator utilities for creating Word documents from markdown outlines.
"""
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_outline_docx(markdown_text, project_id):
    """
    Convert markdown outline to a Word document.
    Returns the path to the generated .docx file.
    """
    # Get project root (3 levels up from src/services/outline_service.py)
    project_root = Path(__file__).parent.parent.parent
    
    # Create output directory if it doesn't exist
    static_dir = project_root / "static"
    static_dir.mkdir(exist_ok=True)
    
    # Output filename
    output_file = static_dir / f"outline_{project_id}.docx"
    
    # Create a new Document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Split content into lines
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        # Skip empty lines
        if not line:
            continue
        
        # H1 heading (# )
        if line.startswith('# '):
            text = line[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
        # H2 heading (## )
        elif line.startswith('## '):
            text = line[3:].strip()
            heading = doc.add_heading(text, level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
        # H3 heading (### )
        elif line.startswith('### '):
            text = line[4:].strip()
            heading = doc.add_heading(text, level=3)
            heading.runs[0].font.color.rgb = RGBColor(51, 51, 51)
            
        # Horizontal rule (---)
        elif line.strip() == '---':
            doc.add_paragraph('_' * 50)
            
        # Bold text (**text**)
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        # Bullet points (- or *)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            # Count indentation level
            indent_level = len(line) - len(line.lstrip())
            text = line.strip()[2:].strip()
            
            # Remove markdown formatting from text
            text = text.replace('**', '')
            
            p = doc.add_paragraph(text, style='List Bullet')
            
            # Adjust indentation
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(indent_level / 8)
        
        # Numbered list
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = text.replace('**', '')
            doc.add_paragraph(text, style='List Number')
        
        # Regular paragraph
        else:
            # Remove markdown formatting
            text = line.replace('**', '')
            if text.strip():
                doc.add_paragraph(text)
    
    # Save the document
    doc.save(str(output_file))
    print(f"✓ Created Word document: {output_file}")
    return str(output_file)


def parse_docx_outline(file_path):
    """
    Parse a document back to markdown text.
    Handles .docx, .odt, .md, and .txt files.
    """
    if file_path.endswith('.md') or  file_path.endswith('.txt'):
        # For markdown files, just read the content
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    elif file_path.endswith('.odt'):
        # For ODT files, extract text from the ZIP archive's content.xml
        # ODT files are ZIP archives containing XML
        try:
            with zipfile.ZipFile(file_path, 'r') as odt_zip:
                # Read the content.xml file which contains the document text
                content_xml = odt_zip.read('content.xml')
            
            # Parse the XML
            root = ET.fromstring(content_xml)
            
            # Define ODF namespaces
            namespaces = {
                'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
                'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0'
            }
            
            # Extract all text content
            text_content = []
            
            # Find all paragraph elements (text:p) and heading elements (text:h)
            for elem in root.iter():
                if elem.tag == '{urn:oasis:names:tc:opendocument:xmlns:text:1.0}h':
                    # Heading element - get outline level for markdown formatting
                    level = int(elem.get('{urn:oasis:names:tc:opendocument:xmlns:text:1.0}outline-level', '1'))
                    text = ''.join(elem.itertext()).strip()
                    if text:
                        text_content.append(f"{'#' * level} {text}")
                        
                elif elem.tag == '{urn:oasis:names:tc:opendocument:xmlns:text:1.0}p':
                    # Regular paragraph
                    text = ''.join(elem.itertext()).strip()
                    if text:
                        text_content.append(text)
                        
                elif elem.tag == '{urn:oasis:names:tc:opendocument:xmlns:text:1.0}list-item':
                    # List item - extract text from child paragraphs
                    for p in elem.findall('.//{urn:oasis:names:tc:opendocument:xmlns:text:1.0}p'):
                        text = ''.join(p.itertext()).strip()
                        if text:
                            text_content.append(f"- {text}")
            
            return '\n'.join(text_content)
            
        except Exception as e:
            raise ValueError(f"Error parsing ODT file: {str(e)}")
    
    elif file_path.endswith('.docx'):
        # For Word documents, extract text and convert to markdown
        doc = Document(file_path)
        markdown_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                markdown_lines.append(f"{'#' * level} {text}")
            
            # Check if it's a list
            elif para.style.name == 'List Bullet':
                markdown_lines.append(f"- {text}")
            elif para.style.name == 'List Number':
                markdown_lines.append(f"1. {text}")
            
            # Regular paragraph
            else:
                markdown_lines.append(text)
        
        return '\n'.join(markdown_lines)
    
    else:
        raise ValueError(f"Unsupported file type: {file_path}")

