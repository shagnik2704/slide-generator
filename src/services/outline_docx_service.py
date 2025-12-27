"""
Course Outline DOCX Export Service.
Converts outline_data JSON to Word document matching the ST Course Outline template.
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_outline_docx(outline_data: dict, output_path: str = None) -> str:
    """
    Generate a DOCX from course outline data matching the ST template.
    
    Args:
        outline_data: The outline_data dict from outline_chat
        output_path: Optional output path. If None, saves to static/
        
    Returns:
        Path to the generated DOCX
    """
    # Get project root
    project_root = Path(__file__).parent.parent.parent
    static_dir = project_root / "static"
    static_dir.mkdir(exist_ok=True)
    
    if output_path is None:
        base_name = outline_data.get('outline_name') or outline_data.get('tutorial_name') or 'outline'
        safe_name = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in base_name)
        output_path = static_dir / f"course_outline_{safe_name[:30]}.docx"
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # === TITLE ===
    title = doc.add_heading('Course Outline - Format', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # === METADATA TABLE ===
    # Determine outline type (FOSS or ICT)
    outline_type = outline_data.get('outline_type', 'FOSS').upper()
    
    # Get platform name (FOSS/ICT tool name and version)
    platform_name = outline_data.get('platform_name', 'Not Applicable for this series')
    if not platform_name or platform_name.strip() == '':
        platform_name = 'Not Applicable for this series'
    
    # Use appropriate label based on outline type
    if outline_type == 'ICT':
        platform_label = "ICT Platform/Program"
    else:
        platform_label = "FOSS Version"
    
    metadata_rows = [
        ("Course Outline Name", outline_data.get('outline_name', outline_data.get('tutorial_name', ''))),
        (platform_label, platform_name),
        ("Target Audience", outline_data.get('target_audience', '')),
        ("Entry Behaviour", outline_data.get('entry_behaviour', '')),
        ("Purpose", outline_data.get('purpose', '')),
        ("OS version", outline_data.get('os_version', 'Not Applicable for this series')),
        ("Recommended no. of tutorials", str(outline_data.get('recommended_no_of_tutorials', ''))),
        ("Prepared by", outline_data.get('prepared_by', '')),
        ("Domain", outline_data.get('domain', '')),
        ("Reviewer", outline_data.get('reviewer', 'IITB ST Team')),
        ("Client Side Reviewer", "Will be from the IITB ST Team. Hence you may leave this blank."),
        ("Date", outline_data.get('date', '')),
        ("Keywords", "; ".join(outline_data.get('keywords', []))),
    ]
    
    table = doc.add_table(rows=len(metadata_rows), cols=2)
    table.style = 'Table Grid'
    
    for i, (label, value) in enumerate(metadata_rows):
        row = table.rows[i]
        # Label cell
        label_cell = row.cells[0]
        label_cell.text = label
        label_cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(label_cell, 'F5F5F5')
        # Value cell
        row.cells[1].text = str(value)
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(5)
    
    doc.add_paragraph()
    
    # === COURSE OBJECTIVES ===
    doc.add_heading('Course Objectives:', level=1)
    for obj in outline_data.get('course_objectives', []):
        p = doc.add_paragraph(obj, style='List Bullet')
    
    # === TOPICS INCLUDED ===
    doc.add_heading('Topics Included', level=1)
    for topic in outline_data.get('topics_included', []):
        doc.add_paragraph(topic, style='List Bullet')
    
    # === TOPICS NOT INCLUDED ===
    doc.add_heading('Topics Not Included', level=1)
    for topic in outline_data.get('topics_not_included', []):
        doc.add_paragraph(topic, style='List Bullet')
    
    # === EXAMPLES TABLE ===
    doc.add_heading('Examples', level=1)
    examples_table = doc.add_table(rows=2, cols=2)
    examples_table.style = 'Table Grid'
    
    # Use appropriate labels based on outline type
    if outline_type == 'ICT':
        core_example_label = "Teaching Scenarios/Examples (core use case)"
        allied_example_label = "Allied examples/scenarios"
    else:
        core_example_label = "Core example used in the series"
        allied_example_label = "Allied examples used in this series"
    
    # Core example
    examples_table.rows[0].cells[0].text = core_example_label
    examples_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_shading(examples_table.rows[0].cells[0], 'F5F5F5')
    examples_table.rows[0].cells[1].text = outline_data.get('core_example', '')
    
    # Allied examples
    examples_table.rows[1].cells[0].text = allied_example_label
    examples_table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_shading(examples_table.rows[1].cells[0], 'F5F5F5')
    examples_table.rows[1].cells[1].text = "; ".join(outline_data.get('allied_examples', []))
    
    doc.add_paragraph()
    
    # === COURSE OUTLINE GUIDELINES ===
    doc.add_heading('Course Outline Guidelines', level=1)
    if outline_type == 'ICT':
        guidelines = [
            "Focus on skill-building and practical application (what learners will DO or TEACH).",
            "Include teaching methodologies and integration strategies.",
            "Use relatable teaching scenarios and real-world educational applications.",
            "Keep content practical and actionable (avoid pure theory).",
            "Organize topics by categories or skill areas when helpful.",
            "Each tutorial should focus on a specific skill, methodology, or integration strategy.",
            "Avoid repetition across tutorials.",
            "Flag topics that are too advanced or off-scope."
        ]
    else:  # FOSS
        guidelines = [
            "Every script should be written such that 75-80% of the content is the demonstration.",
            "Keep theory to a minimum.",
            "Do not use the Menu-based approach to explain the features of the software.",
            "Cover crucial and important features in detail.",
            "Avoid repetitions across the outline or across scripts or within a script.",
            "Provide quick pointers for less important features.",
            "Avoid spending too much time on topics which learners can figure out on their own.",
            "Refer to other features in assignments with an appropriate number of hints to help the learner."
        ]
    for g in guidelines:
        doc.add_paragraph(g, style='List Bullet')
    
    doc.add_paragraph()
    
    # === TUTORIAL TABLES ===
    for tutorial in outline_data.get('tutorial_rows', []):
        tutorial_num = tutorial.get('tutorial_number', 1)
        title_text = tutorial.get('title', f'Tutorial {tutorial_num}')
        
        doc.add_heading(f'Tutorial Title {tutorial_num}: {title_text}', level=2)
        
        # Handle prerequisites as list or string (for backward compatibility)
        prerequisites_data = tutorial.get('prerequisites', [])
        if isinstance(prerequisites_data, list):
            prerequisites = '; '.join(prerequisites_data) if prerequisites_data else 'N/A'
        else:
            prerequisites = prerequisites_data if prerequisites_data else 'N/A'
        
        # Create tutorial table
        topics = tutorial.get('topics_details', [])
        num_rows = max(len(topics), 4) + 1  # At least 4 rows + header
        
        tutorial_table = doc.add_table(rows=num_rows, cols=4)
        tutorial_table.style = 'Table Grid'
        
        # Header row
        header_row = tutorial_table.rows[0]
        headers = ["Prerequisites", "Topics Details", "Time (secs)", "Comments"]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'E8E8E8')
        
        # Data rows
        time_seconds = tutorial.get('time_seconds', 180)
        comments = tutorial.get('comments', '')
        
        # Display time as range if available
        time_range = tutorial.get('time_range')
        if time_range:
            min_minutes = time_range.get('min_seconds', 0) // 60
            max_minutes = time_range.get('max_seconds', 0) // 60
            if min_minutes == max_minutes:
                time_display = f"{min_minutes} min"
            else:
                time_display = f"{min_minutes}-{max_minutes} min"
        else:
            # Fallback to time_seconds
            time_display = f"{time_seconds // 60} min" if time_seconds > 0 else "0 min"
        
        for i, topic in enumerate(topics):
            row = tutorial_table.rows[i + 1]
            row.cells[0].text = prerequisites if i == 0 else ""
            row.cells[1].text = f"{i + 1}. {topic}"
            row.cells[2].text = time_display if i == 0 else ""
            row.cells[3].text = comments if i == 0 else ""
        
        # Empty rows to match template
        for i in range(len(topics) + 1, num_rows):
            row = tutorial_table.rows[i]
            row.cells[0].text = ""
            row.cells[1].text = f"{i}. "
            row.cells[2].text = ""
            row.cells[3].text = ""
        
        doc.add_paragraph()
    
    # Save document
    doc.save(str(output_path))
    print(f"✅ Generated outline DOCX: {output_path}")
    
    return str(output_path)


if __name__ == "__main__":
    # Test with sample data
    sample_data = {
        "tutorial_name": "Ethics, Responsibility and Future of AI",
        "target_audience": "Everyone",
        "entry_behaviour": "Basics of generative AI",
        "purpose": "Understand the societal impact of GenAI",
        "recommended_no_of_tutorials": 1,
        "prepared_by": "Test User",
        "date": "2024-12-18",
        "keywords": ["GenAI", "AI", "Ethics"],
        "about_course": "This course teaches about AI ethics and responsibility.",
        "course_objectives": ["Understand AI ethics", "Recognize misinformation"],
        "topics_included": ["Bias in AI", "Deepfakes"],
        "topics_not_included": ["Coding"],
        "core_example": "An AI trained on biased data...",
        "allied_examples": ["Deepfake videos", "Fake audio"],
        "tutorial_rows": [
            {
                "tutorial_number": 1,
                "title": "Introduction to AI Ethics",
                "topics_details": ["What is AI bias", "Examples of bias", "How to detect bias"],
                "time_seconds": 300,
                "comments": ""
            }
        ]
    }
    create_outline_docx(sample_data)
