"""Timed Script route - Generate sentence-level timestamps from audio."""

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Body
from fastapi.encoders import jsonable_encoder
from pathlib import Path
from uuid import UUID, uuid4

from src.api.auth import get_current_user, TokenData
from src.jobs.persistence import (
    attach_celery_task,
    create_job,
    fail_job,
    get_job,
    list_jobs,
)
from src.workers.celery_app import celery_app

router = APIRouter(prefix="/timed-script", tags=["timed-script"])

TIMED_SCRIPT_UPLOAD_DIR = Path(__file__).resolve().parents[3] / "uploads" / "timed_script_jobs"


def _public_job(job: dict) -> dict:
    return jsonable_encoder(
        {
            "job_id": str(job["id"]),
            "job_type": job["job_type"],
            "status": job["status"],
            "original_filename": job["original_filename"],
            "result": job["result"],
            "error_message": job["error_message"],
            "progress": job["progress"],
            "current_stage": job["current_stage"],
            "created_at": job["created_at"],
            "started_at": job["started_at"],
            "completed_at": job["completed_at"],
            "updated_at": job["updated_at"],
        }
    )


@router.post("/generate")
async def generate_timed_script_endpoint(
    audio: UploadFile = File(...),
    language: str = None,
    current_user: TokenData = Depends(get_current_user)
):
    """
    Generate a timed script from an uploaded audio file.
    
    Returns sentence-level timestamps with time ranges.
    Language is auto-detected by default.
    
    Args:
        audio: Audio/Video file (WAV, MP3, MP4, etc.)
        language: Language code (optional, auto-detect if not provided)
        
    Returns a queued job. The Whisper worker performs transcription after the
    request returns, so the browser does not need to remain connected.
    """
    allowed_extensions = {'.wav', '.mp3', '.m4a', '.ogg', '.flac', '.webm', '.mp4'}
    file_ext = Path(audio.filename or "").suffix.lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Allowed: {', '.join(allowed_extensions)}"
        )
    
    # Keep the first implementation on the shared uploads volume. Object
    # storage can replace this handoff later without changing the job contract.
    target_path = TIMED_SCRIPT_UPLOAD_DIR / f"{uuid4().hex}{file_ext}"
    try:
        TIMED_SCRIPT_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
        with target_path.open("wb") as target:
            while chunk := await audio.read(1024 * 1024):
                target.write(chunk)

        job = await create_job(
            user_id=current_user.sub,
            job_type="timed_script",
            input_path=str(target_path),
            original_filename=audio.filename,
        )
        try:
            celery_task = celery_app.send_task(
                "src.workers.tasks.process_timed_script",
                args=[str(job["id"]), language],
            )
            await attach_celery_task(str(job["id"]), celery_task.id)
        except Exception as enqueue_error:
            await fail_job(str(job["id"]), f"Could not queue job: {enqueue_error}")
            target_path.unlink(missing_ok=True)
            raise HTTPException(status_code=503, detail="Background worker is unavailable") from enqueue_error

        return _public_job(job)
    except Exception as e:
        target_path.unlink(missing_ok=True)
        if isinstance(e, HTTPException):
            raise
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/jobs")
async def get_timed_script_jobs(current_user: TokenData = Depends(get_current_user)):
    """List the current user's timed-script jobs."""
    jobs = await list_jobs(current_user.sub, job_type="timed_script")
    return {"jobs": [_public_job(job) for job in jobs]}


@router.get("/jobs/{job_id}")
async def get_timed_script_job(job_id: str, current_user: TokenData = Depends(get_current_user)):
    """Return one timed-script job without exposing its input path."""
    # job_id maps to a uuid column; reject malformed ids as 404 rather than
    # letting the invalid-uuid cast surface as a 500.
    try:
        UUID(job_id)
    except (TypeError, ValueError):
        raise HTTPException(status_code=404, detail="Timed script job not found")
    job = await get_job(job_id, current_user.sub)
    if not job or job["job_type"] != "timed_script":
        raise HTTPException(status_code=404, detail="Timed script job not found")
    return _public_job(job)


@router.post("/download-docx")
async def download_timed_script_docx(
    data: dict = Body(...),
    current_user: TokenData = Depends(get_current_user)
):
    """
    Generate a DOCX file from timed script results.
    
    Args:
        data: The result from /generate endpoint containing sentences
        
    Returns:
        DOCX file download
    """
    from fastapi.responses import StreamingResponse
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    sentences = data.get("sentences", [])
    audio_file = data.get("audio_file", "audio")
    total_duration = data.get("total_duration", "00:00")
    
    if not sentences:
        raise HTTPException(status_code=400, detail="No sentences provided")
    
    doc = Document()
    
    title = doc.add_heading("Timed Script", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Audio File: {audio_file}")
    doc.add_paragraph(f"Total Duration: {total_duration}")
    doc.add_paragraph(f"Total Sentences: {len(sentences)}")
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Time Range"
    hdr_cells[1].text = "Text"
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for sentence in sentences:
        row_cells = table.add_row().cells
        row_cells[0].text = sentence.get("time_range", "")
        row_cells[1].text = sentence.get("text", "")
    
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.0)
        
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                '<w:top w:w="200" w:type="dxa"/>'
                '<w:left w:w="200" w:type="dxa"/>'
                '<w:bottom w:w="200" w:type="dxa"/>'
                '<w:right w:w="200" w:type="dxa"/>'
                '</w:tcMar>'
            )
            tcPr.append(tcMar)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    base_name = Path(audio_file).stem if audio_file else "timed_script"
    filename = f"{base_name}_timed_script.docx"
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
