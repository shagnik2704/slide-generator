"""Compliance check route handlers."""
from fastapi import APIRouter, HTTPException, File, UploadFile, Depends
from pathlib import Path
import os
import json
import time
import traceback

from src.api.auth import get_current_user, TokenData

router = APIRouter(tags=["compliance"])


def _detect_tutorial_type(json_script: dict) -> str:
    """Detect if script is demo or conceptual based on content."""
    slides = json_script.get('slides', [])
    
    # Check for demo-like content (action verbs, step-by-step)
    action_verbs = ['click', 'open', 'type', 'select', 'navigate', 'copy', 'paste']
    action_count = 0
    
    for slide in slides:
        narration = slide.get('narration', '').lower()
        if any(verb in narration for verb in action_verbs):
            action_count += 1
    
    # If more than half the slides have action verbs, it's a demo
    if action_count > len(slides) / 2:
        return "demo"
    return "conceptual"


@router.post("/check_compliance")
async def check_compliance_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Run compliance checks on a parsed script.
    
    Args:
        json_script: The parsed script JSON
        tutorial_type: Optional, 'demo' or 'conceptual' (auto-detected if not provided)
    
    Returns:
        Compliance report with checks and summary
    """
    print("Running compliance checks...")
    
    try:
        json_script = data.get('json_script')
        
        if not json_script:
            raise HTTPException(status_code=400, detail="json_script is required")
        
        # Get tutorial type or detect it
        tutorial_type = data.get('tutorial_type')
        if not tutorial_type:
            tutorial_type = _detect_tutorial_type(json_script)
        
        from src.services.compliance_service import check_compliance
        compliance_report = await check_compliance(json_script, tutorial_type)
        
        summary = compliance_report.get('summary', {})
        print(f"✅ Compliance check complete: {summary.get('ai_passed', 0)} passed, {summary.get('ai_failed', 0)} failed")
        
        return compliance_report
        
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in check_compliance: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/check_admin_compliance_v1")
async def check_admin_compliance_v1_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Run evidence-based admin compliance checks on a parsed script.

    This is the parallel V1 workflow. It keeps legacy-friendly checks/summary
    while adding issues, annotations, policy, and artifact_summary.
    """
    print("Running evidence-based admin compliance checks...")

    try:
        json_script = data.get('json_script')
        if not json_script:
            raise HTTPException(status_code=400, detail="json_script is required")

        tutorial_type = data.get('tutorial_type') or _detect_tutorial_type(json_script)
        source_artifact = data.get('source_artifact') or {}

        from src.compliance.workflow import run_admin_script_compliance
        compliance_report = await run_admin_script_compliance(
            json_script,
            tutorial_type=tutorial_type,
            source_artifact=source_artifact,
        )

        summary = compliance_report.get('summary', {})
        print(
            "✅ Evidence-based compliance complete: "
            f"{summary.get('ai_passed', 0)} passed, "
            f"{summary.get('ai_failed', 0)} failed, "
            f"{summary.get('ai_skipped', 0)} skipped"
        )

        return compliance_report

    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in check_admin_compliance_v1: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/check_outline_compliance")
async def check_outline_compliance_endpoint(data: dict):
    """
    Run compliance checks on an outline/tutorial design.
    
    Args:
        outline_data: The outline JSON data (CourseOutlineData format)
    
    Returns:
        Compliance report with checks and summary for outline design
    """
    print("Running outline compliance checks...")
    
    try:
        outline_data = data.get('outline_data')
        
        if not outline_data:
            raise HTTPException(status_code=400, detail="outline_data is required")
        
        from src.services.compliance_service import check_outline_compliance
        compliance_report = await check_outline_compliance(outline_data)
        
        summary = compliance_report.get('summary', {})
        print(f"✅ Outline compliance check complete: {summary.get('ai_passed', 0)} passed, {summary.get('ai_failed', 0)} failed")
        
        return compliance_report
        
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in check_outline_compliance: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload_outline_for_compliance")
async def upload_outline_for_compliance(file: UploadFile = File(...)):
    """
    Upload an outline file (.json or .docx) and run compliance checks on it.
    
    Args:
        file: JSON file containing outline_data (CourseOutlineData format) OR DOCX file to parse
    
    Returns:
        Compliance report with checks and summary for outline design
    """
    print(f"Uploading outline file for compliance check: {file.filename}")
    
    try:
        filename = file.filename.lower()
        
        # Validate file type - support JSON and DOCX files
        if not (filename.endswith('.json') or filename.endswith('.docx')):
            raise HTTPException(status_code=400, detail="Only .json or .docx files are allowed. JSON files should contain 'outline_data' field. DOCX files will be parsed to extract outline data.")
        
        # Read file content
        content = await file.read()
        
        # Handle DOCX files
        if filename.endswith('.docx'):
            # Parse DOCX and extract outline data using LLM
            from src.services.outline_service import parse_docx_outline
            
            # Save temporarily to parse
            project_root = Path(__file__).parent.parent.parent
            upload_dir = project_root / "uploads"
            upload_dir.mkdir(exist_ok=True)
            temp_path = upload_dir / f"outline_compliance_{int(time.time())}_{file.filename}"
            
            with open(str(temp_path), "wb") as buffer:
                buffer.write(content)
            
            try:
                # Parse DOCX to markdown text
                outline_text = parse_docx_outline(str(temp_path))
                
                # Use LLM to extract outline_data from the parsed text
                from src.api.routes.outline_chat.outline_chat_llm_utils import generate_llm_text
                import asyncio
                
                # Create a prompt to extract outline data from the text
                extraction_prompt = f"""Extract course outline data from the following document text and return it as a JSON object matching the CourseOutlineData format.

Document text:
{outline_text}

Return a JSON object with the following structure (fill in what you can find):
{{
    "outline_type": "FOSS" or "ICT" or "OTHER",
    "outline_name": "course name",
    "platform_name": "software/platform name and version",
    "target_audience": "target audience description",
    "entry_behaviour": "prerequisites/entry behaviour",
    "purpose": "course purpose",
    "os_version": "OS version if applicable",
    "recommended_no_of_tutorials": number,
    "prepared_by": "author name",
    "domain": "domain name",
    "reviewer": "reviewer name",
    "date": "date",
    "keywords": ["keyword1", "keyword2"],
    "course_objectives": ["objective1", "objective2"],
    "topics_included": ["topic1", "topic2"],
    "topics_not_included": ["topic1", "topic2"],
    "core_example": "core example description",
    "allied_examples": ["example1", "example2"],
    "tutorial_rows": [
        {{
            "tutorial_number": 1,
            "title": "tutorial title",
            "prerequisites": ["prereq1"],
            "topics_details": ["topic1", "topic2"],
            "time_seconds": 180,
            "comments": "comments"
        }}
    ]
}}

Return ONLY the JSON object, no other text."""
                
                # Call LLM to extract outline data (run in thread to avoid blocking)
                llm_response = await asyncio.to_thread(
                    generate_llm_text,
                    extraction_prompt,
                    temperature=0.3,
                    max_tokens=4096
                )
                
                # Parse the LLM response to extract JSON
                import re
                json_match = re.search(r'\{.*\}', llm_response, re.DOTALL)
                if json_match:
                    outline_data = json.loads(json_match.group())
                else:
                    # Try to parse the entire response as JSON
                    outline_data = json.loads(llm_response)
                
            except json.JSONDecodeError as e:
                if temp_path.exists():
                    os.remove(str(temp_path))
                raise HTTPException(status_code=400, detail=f"Failed to extract outline data from DOCX. The document may not be in the expected format. Error: {str(e)}")
            except Exception as parse_error:
                if temp_path.exists():
                    os.remove(str(temp_path))
                raise HTTPException(status_code=400, detail=f"Failed to parse DOCX file: {str(parse_error)}")
            finally:
                # Clean up temp file
                if temp_path.exists():
                    os.remove(str(temp_path))
        else:
            # Handle JSON files
            try:
                file_data = json.loads(content.decode('utf-8'))
            except json.JSONDecodeError as e:
                raise HTTPException(status_code=400, detail=f"Invalid JSON format: {str(e)}")
            
            # Extract outline_data - support both direct outline_data or wrapped in a dict
            outline_data = file_data.get('outline_data') or file_data
            
            if not outline_data or not isinstance(outline_data, dict):
                raise HTTPException(status_code=400, detail="JSON file must contain 'outline_data' field or be a valid outline_data object")
        
        # Run compliance checks
        from src.services.compliance_service import check_outline_compliance
        compliance_report = await check_outline_compliance(outline_data)
        
        summary = compliance_report.get('summary', {})
        print(f"✅ Outline compliance check complete for {file.filename}: {summary.get('ai_passed', 0)} passed, {summary.get('ai_failed', 0)} failed")
        
        return {
            "compliance_report": compliance_report,
            "outline_data": outline_data,
            "message": f"Compliance check complete for {file.filename}"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in upload_outline_for_compliance: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/batch_check_compliance")
async def batch_check_compliance_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Check multiple scripts for compliance in parallel.
    
    Args (in data):
        scripts: List of JSON scripts to check
        tutorial_types: Optional list of tutorial types (one per script)
    
    Returns:
        results: List of compliance check results, one per script
        summary: Overall batch summary
    
    Example request:
        {
            "scripts": [script1_json, script2_json, script3_json],
            "tutorial_types": ["conceptual", "demo", "conceptual"]  // optional
        }
    """
    try:
        scripts = data.get('scripts', [])
        tutorial_types = data.get('tutorial_types')
        
        if not scripts:
            raise HTTPException(status_code=400, detail="No scripts provided")
        
        if not isinstance(scripts, list):
            raise HTTPException(status_code=400, detail="scripts must be a list")
        
        print(f"📋 Batch compliance check: {len(scripts)} scripts")
        
        from src.services.compliance_service import batch_check_compliance
        results = await batch_check_compliance(scripts, tutorial_types)
        
        # Calculate overall summary
        total_passed = sum(
            1 for r in results 
            if r.get('summary', {}).get('ai_failed', 1) == 0
        )
        
        return {
            "results": results,
            "batch_summary": {
                "total_scripts": len(scripts),
                "scripts_passed": total_passed,
                "scripts_with_issues": len(scripts) - total_passed
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in batch_check_compliance: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/export_compliance_report")
async def export_compliance_report(data: dict, current_user: TokenData = Depends(get_current_user)):
    """Export compliance report as DOCX or ODT file."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from io import BytesIO
    from fastapi.responses import StreamingResponse
    
    try:
        checks = data.get('checks', [])
        summary = data.get('summary', {})
        format_type = data.get('format', 'docx')  # 'docx' or 'odt'
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Script Compliance Report', level=0)
        
        # Add summary
        doc.add_paragraph(
            f"AI Passed: {summary.get('ai_passed', 0)} | "
            f"AI Failed: {summary.get('ai_failed', 0)} | "
            f"Total: {summary.get('total', 0)}"
        )
        doc.add_paragraph()
        
        # Create table - Same format as inline: # | Criteria | AI | AI Notes | Human Review
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = '#'
        header_cells[1].text = 'Criteria'
        header_cells[2].text = 'AI'
        header_cells[3].text = 'AI Notes'
        header_cells[4].text = 'Human Review'
        
        # Set column widths
        from docx.shared import Cm
        table.columns[0].width = Cm(1)      # # - narrow
        table.columns[1].width = Cm(6)      # Criteria - wide
        table.columns[2].width = Cm(1)      # AI - narrow
        table.columns[3].width = Cm(5)      # AI Notes - wide
        table.columns[4].width = Cm(3)      # Human Review
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for i, check in enumerate(checks):
            row = table.add_row()
            
            # Row number
            row.cells[0].text = str(i + 1)
            
            # Criteria
            row.cells[1].text = check.get('criteria', '')
            
            # AI Status - Tick/Cross
            ai_review = check.get('ai_review')
            if ai_review is True:
                row.cells[2].text = '✓'
            elif ai_review is False:
                row.cells[2].text = '✗'
            else:
                row.cells[2].text = '—'
            
            # AI Notes
            row.cells[3].text = check.get('ai_notes', '')
            
            # Human Review (editable text from user)
            human_review = check.get('human_review', '')
            row.cells[4].text = str(human_review) if human_review else ''
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Return file
        filename = f"compliance_report.{format_type}"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        return StreamingResponse(
            buffer,
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/export_admin_compliance_review")
async def export_admin_compliance_review(data: dict, current_user: TokenData = Depends(get_current_user)):
    """Export script DOCX in the standard script format with review columns."""
    from io import BytesIO
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    from fastapi.responses import StreamingResponse
    from src.services.docx_service import (
        _add_formatted_text,
        _add_metadata_section,
        _format_visual_cue,
        _set_cell_padding,
    )

    def row_id_for_index(index: int) -> str:
        return f"row_{index + 1:03d}"

    def issue_summary(issue: dict, check_by_id: dict) -> str:
        check = check_by_id.get(issue.get("criteria_id"), {})
        criterion = check.get("criteria") or issue.get("criteria_id") or "Compliance issue"
        message = issue.get("message") or check.get("ai_notes") or ""
        severity = issue.get("severity") or "issue"
        return f"[{severity.upper()}] {criterion}" + (f"\n{message}" if message else "")

    try:
        json_script = data.get("json_script") or {}
        report = data.get("report") or {}
        comments = data.get("review_comments") or {}
        slides = json_script.get("slides") or []
        checks = report.get("checks") or []
        issues = report.get("issues") or []
        summary = report.get("summary") or {}

        check_by_id = {check.get("id"): check for check in checks if check.get("id")}
        issues_by_row = {}
        for issue in issues:
            for evidence in issue.get("evidence") or []:
                row_id = evidence.get("row_id")
                if row_id:
                    issues_by_row.setdefault(row_id, []).append(issue_summary(issue, check_by_id))

        title = json_script.get("presentation_title", "Spoken Tutorial Script")
        doc = Document()
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _add_metadata_section(doc, json_script)

        doc.add_heading("Script", level=1)
        doc.add_paragraph(
            "Review the AI feedback and add reviewer notes in the final column. "
            "Do not modify the table structure."
        )

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Inches(2.1)
        table.columns[1].width = Inches(3.0)
        table.columns[2].width = Inches(2.0)
        table.columns[3].width = Inches(2.0)

        header_cells = table.rows[0].cells
        header_cells[0].text = "Visual Cue"
        header_cells[1].text = "Narration"
        header_cells[2].text = "AI Review"
        header_cells[3].text = "Human Notes"
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(12)

        for index, slide in enumerate(slides):
            row_id = slide.get("row_id") or row_id_for_index(index)
            row = table.add_row()
            _add_formatted_text(row.cells[0], _format_visual_cue(slide))
            narration = str(slide.get("narration") or "").replace("\\n", "\n")
            _add_formatted_text(row.cells[1], narration)
            _add_formatted_text(row.cells[2], "\n\n".join(issues_by_row.get(row_id, [])))
            _add_formatted_text(row.cells[3], str(comments.get(row_id) or ""))

            for cell in row.cells:
                _set_cell_padding(cell, top=100, bottom=100, left=100, right=100)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=admin_compliance_script_review.docx"},
        )

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
