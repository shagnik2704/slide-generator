"""Upload route handlers."""
from fastapi import APIRouter, HTTPException, File, UploadFile, Form, Depends
from pathlib import Path
import os
import json
import time
import traceback

from src.api.auth import get_current_user, TokenData
from src.services.outline_service import parse_docx_outline

router = APIRouter(tags=["upload"])


@router.post("/upload_outline")
async def upload_outline(file: UploadFile = File(...), current_user: TokenData = Depends(get_current_user)):
    """Upload an edited outline file (.md or .docx)."""
    try:
        # Validate file type
        if not (file.filename.endswith('.md') or file.filename.endswith('.docx') or file.filename.endswith('.txt') or file.filename.endswith('.odt')):
            raise HTTPException(status_code=400, detail="Only .md, .txt, .docx, or .odt files are allowed")
        
        # Get project root (3 levels up from src/api/routes/upload.py)
        project_root = Path(__file__).parent.parent.parent
        
        # Save uploaded file temporarily
        upload_dir = project_root / "uploads"
        upload_dir.mkdir(exist_ok=True)
        
        temp_path = upload_dir / f"outline_{int(time.time())}_{file.filename}"
        
        with open(str(temp_path), "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Parse the document
        outline_text = parse_docx_outline(str(temp_path))
        
        # Clean up temp file
        os.remove(str(temp_path))
        
        return {
            "outline": outline_text,
            "message": "Outline uploaded successfully"
        }
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in upload_outline: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/parse_script")
async def parse_script(file: UploadFile = File(...), current_user: TokenData = Depends(get_current_user)):
    """Parse a script file (.json, .docx, or .odt) WITHOUT running any checks."""
    try:
        filename = file.filename.lower()
        
        # Validate file type
        if not (filename.endswith('.json') or filename.endswith('.docx') or filename.endswith('.odt')):
            raise HTTPException(status_code=400, detail="Only .json, .docx, or .odt files are allowed")
        
        # Get project root
        project_root = Path(__file__).parent.parent.parent
        
        # Read file content
        content = await file.read()
        
        # Parse based on file type
        if filename.endswith('.json'):
            # Direct JSON parsing
            json_script = json.loads(content.decode('utf-8'))
        else:
            # Parse docx/odt using existing parser
            from io import BytesIO
            from src.services.docx_service import docx_to_json
            
            json_script = docx_to_json(BytesIO(content))
        
        # Detect tutorial type (demo or conceptual)
        tutorial_type = _detect_tutorial_type(json_script)
        
        # Generate project ID
        project_id = int(time.time())
        
        # Save a copy
        output_dir = project_root / "output"
        output_dir.mkdir(exist_ok=True)
        json_path = output_dir / f"script_{project_id}.json"
        with open(str(json_path), 'w') as f:
            json.dump(json_script, f, indent=2)
        
        print(f"✅ Script parsed: {file.filename} → project #{project_id}")
        
        return {
            "json_script": json_script,
            "project_id": project_id,
            "tutorial_type": tutorial_type,
            "message": "Script parsed successfully"
        }
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON format: {str(e)}")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=f"Could not parse document: {str(e)}")
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in parse_script: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/check_compliance")
async def check_compliance_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Run compliance checks on a parsed script.
    
    Args:
        json_script: The parsed script JSON
        tutorial_type: Optional, 'demo' or 'conceptual' (auto-detected if not provided)
    
    Returns:
        Compliance report with checks and summary
    """
    print("Running compliance checks...")
    
    try:
        json_script = data.get('json_script')
        
        if not json_script:
            raise HTTPException(status_code=400, detail="json_script is required")
        
        # Get tutorial type or detect it
        tutorial_type = data.get('tutorial_type')
        if not tutorial_type:
            tutorial_type = _detect_tutorial_type(json_script)
        
        from src.services.compliance_service import check_compliance
        compliance_report = await check_compliance(json_script, tutorial_type)
        
        summary = compliance_report.get('summary', {})
        print(f"✅ Compliance check complete: {summary.get('ai_passed', 0)} passed, {summary.get('ai_failed', 0)} failed")
        
        return compliance_report
        
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in check_compliance: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/batch_check_compliance")
async def batch_check_compliance_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Check multiple scripts for compliance in parallel.
    
    Args (in data):
        scripts: List of JSON scripts to check
        tutorial_types: Optional list of tutorial types (one per script)
    
    Returns:
        results: List of compliance check results, one per script
        summary: Overall batch summary
    
    Example request:
        {
            "scripts": [script1_json, script2_json, script3_json],
            "tutorial_types": ["conceptual", "demo", "conceptual"]  // optional
        }
    """
    try:
        scripts = data.get('scripts', [])
        tutorial_types = data.get('tutorial_types')
        
        if not scripts:
            raise HTTPException(status_code=400, detail="No scripts provided")
        
        if not isinstance(scripts, list):
            raise HTTPException(status_code=400, detail="scripts must be a list")
        
        print(f"📋 Batch compliance check: {len(scripts)} scripts")
        
        from src.services.compliance_service import batch_check_compliance
        results = await batch_check_compliance(scripts, tutorial_types)
        
        # Calculate overall summary
        total_passed = sum(
            1 for r in results 
            if r.get('summary', {}).get('ai_failed', 1) == 0
        )
        
        return {
            "results": results,
            "batch_summary": {
                "total_scripts": len(scripts),
                "scripts_passed": total_passed,
                "scripts_with_issues": len(scripts) - total_passed
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in batch_check_compliance: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload_script")
async def upload_script(file: UploadFile = File(...), current_user: TokenData = Depends(get_current_user)):
    """Upload a script file (.json, .docx, or .odt) and run compliance check."""
    try:
        filename = file.filename.lower()
        
        # Validate file type
        if not (filename.endswith('.json') or filename.endswith('.docx') or filename.endswith('.odt')):
            raise HTTPException(status_code=400, detail="Only .json, .docx, or .odt files are allowed")
        
        # Get project root
        project_root = Path(__file__).parent.parent.parent
        
        # Read file content
        content = await file.read()
        
        # Parse based on file type
        if filename.endswith('.json'):
            # Direct JSON parsing
            json_script = json.loads(content.decode('utf-8'))
        else:
            # Parse docx/odt using existing parser
            from io import BytesIO
            from src.services.docx_service import docx_to_json
            
            json_script = docx_to_json(BytesIO(content))
        
        # Detect tutorial type (demo or conceptual)
        tutorial_type = _detect_tutorial_type(json_script)
        
        # Run compliance checks
        from src.services.compliance_service import check_compliance
        compliance_report = await check_compliance(json_script, tutorial_type)
        
        # Generate project ID
        project_id = int(time.time())
        
        # Save a copy
        output_dir = project_root / "output"
        output_dir.mkdir(exist_ok=True)
        json_path = output_dir / f"script_{project_id}.json"
        with open(str(json_path), 'w') as f:
            json.dump(json_script, f, indent=2)
        
        print(f"✅ Script uploaded: {file.filename} → project #{project_id}")
        print(f"📋 Compliance: {compliance_report['summary']['ai_failed']} issues found")
        
        return {
            "json_script": json_script,
            "project_id": project_id,
            "message": "Script uploaded successfully",
            "compliance_report": compliance_report
        }
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON format: {str(e)}")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=f"Could not parse document: {str(e)}")
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in upload_script: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _detect_tutorial_type(json_script: dict) -> str:
    """Detect if script is demo or conceptual based on content."""
    slides = json_script.get('slides', [])
    
    # Check for demo-like content (action verbs, step-by-step)
    action_verbs = ['click', 'open', 'type', 'select', 'navigate', 'copy', 'paste']
    action_count = 0
    
    for slide in slides:
        narration = slide.get('narration', '').lower()
        if any(verb in narration for verb in action_verbs):
            action_count += 1
    
    # If more than half the slides have action verbs, it's a demo
    if action_count > len(slides) / 2:
        return "demo"
    return "conceptual"


@router.post("/export_compliance_report")
async def export_compliance_report(data: dict, current_user: TokenData = Depends(get_current_user)):
    """Export compliance report as DOCX or ODT file."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from io import BytesIO
    from fastapi.responses import StreamingResponse
    
    try:
        checks = data.get('checks', [])
        summary = data.get('summary', {})
        format_type = data.get('format', 'docx')  # 'docx' or 'odt'
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Script Compliance Report', level=0)
        
        # Add summary
        doc.add_paragraph(
            f"AI Passed: {summary.get('ai_passed', 0)} | "
            f"AI Failed: {summary.get('ai_failed', 0)} | "
            f"Total: {summary.get('total', 0)}"
        )
        doc.add_paragraph()
        
        # Create table - Same format as inline: # | Criteria | AI | AI Notes | Human Review
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = '#'
        header_cells[1].text = 'Criteria'
        header_cells[2].text = 'AI'
        header_cells[3].text = 'AI Notes'
        header_cells[4].text = 'Human Review'
        
        # Set column widths
        from docx.shared import Cm
        table.columns[0].width = Cm(1)      # # - narrow
        table.columns[1].width = Cm(6)      # Criteria - wide
        table.columns[2].width = Cm(1)      # AI - narrow
        table.columns[3].width = Cm(5)      # AI Notes - wide
        table.columns[4].width = Cm(3)      # Human Review
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for i, check in enumerate(checks):
            row = table.add_row()
            
            # Row number
            row.cells[0].text = str(i + 1)
            
            # Criteria
            row.cells[1].text = check.get('criteria', '')
            
            # AI Status - Tick/Cross
            ai_review = check.get('ai_review')
            if ai_review is True:
                row.cells[2].text = '✓'
            elif ai_review is False:
                row.cells[2].text = '✗'
            else:
                row.cells[2].text = '—'
            
            # AI Notes
            row.cells[3].text = check.get('ai_notes', '')
            
            # Human Review (editable text from user)
            human_review = check.get('human_review', '')
            row.cells[4].text = str(human_review) if human_review else ''
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Return file
        filename = f"compliance_report.{format_type}"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        return StreamingResponse(
            buffer,
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/check_quality")
async def check_quality_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Run quality checks and translate script to Hindi.
    
    Returns:
        - Quality check results (translation quality, timing, transliteration)
        - Full translated Hindi script
    """
    print("Running quality checks and Hindi translation...")
    
    try:
        json_script = data.get('json_script')
        
        if not json_script:
            raise HTTPException(status_code=400, detail="json_script is required")
        
        from src.services.quality_service import check_quality
        result = await check_quality(json_script)
        
        summary = result.get('summary', {})
        print(f"✅ Quality check complete: {summary.get('ai_passed', 0)}/{summary.get('total', 0)} passed")
        print(f"📊 Avg translation quality: {summary.get('avg_quality_score', 0)}/5")
        
        return result
        
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in check_quality: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/batch_check_quality")
async def batch_check_quality_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Run quality checks for multiple scripts in parallel.
    
    Args:
        scripts: List of script JSON objects to check
    
    Returns:
        results: List of quality check results (one per script)
        batch_summary: Overall batch statistics
    """
    print("📋 Batch quality check requested...")
    
    try:
        scripts = data.get('scripts', [])
        
        if not scripts:
            raise HTTPException(status_code=400, detail="scripts list is required")
        
        print(f"   Processing {len(scripts)} scripts...")
        
        from src.services.quality_service import batch_check_quality
        result = await batch_check_quality(scripts)
        
        summary = result.get('batch_summary', {})
        print(f"✅ Batch quality check complete: {summary.get('scripts_passed', 0)}/{summary.get('total_scripts', 0)} passed")
        
        return result
        
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in batch_check_quality: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate_voice")
async def generate_voice_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Generate voice narration for a JSON script.
    
    Args:
        json_script: The parsed script JSON
        project_id: Optional project ID (auto-generated if not provided)
    
    Returns:
        audio_urls: Per-slide audio file URLs
        zip_url: URL to download all audio as ZIP
    """
    print("🎤 Starting voice generation...")
    
    try:
        json_script = data.get('json_script')
        project_id = data.get('project_id')
        
        if not json_script:
            raise HTTPException(status_code=400, detail="json_script is required")
        
        from src.services.voice_service import generate_voice_for_script
        result = await generate_voice_for_script(
            json_script=json_script,
            project_id=project_id,
        )
        
        print(f"✅ Voice generation complete: {result.get('generated_slides')}/{result.get('total_slides')} slides")
        
        return result
        
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in generate_voice: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate_voice_batched")
async def generate_voice_batched_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Generate voice narration using BATCHED approach - ~75% fewer API calls!
    
    Combines multiple slides per API call, then splits audio using silence detection.
    
    Args:
        json_script: The parsed script JSON
        project_id: Optional project ID
    """
    print("🎤 Starting BATCHED voice generation...")
    
    try:
        json_script = data.get('json_script')
        project_id = data.get('project_id')
        
        if not json_script:
            raise HTTPException(status_code=400, detail="json_script is required")
        
        from src.services.voice_service import generate_voice_for_script_batched
        result = await generate_voice_for_script_batched(
            json_script=json_script,
            project_id=project_id,
        )
        
        print(f"✅ Batched voice generation complete: {result.get('generated_slides')}/{result.get('total_slides')} slides")
        print(f"📊 API calls saved: {result.get('api_calls_saved')}")
        
        return result
        
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in generate_voice_batched: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/generate_voice_combined")
async def generate_voice_combined_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Generate a SINGLE audio file for the entire script.
    
    Combines all slide narrations with pauses into one continuous audio file.
    Useful for creating full tutorial audio in one go.
    
    Args:
        json_script: The parsed script JSON
        project_id: Optional project ID (auto-generated if not provided)
    
    Returns:
        audio_url: URL to the combined audio file
        duration_estimate: Approximate duration
        total_slides: Number of slides combined
    """
    print("🎤 Starting COMBINED voice generation...")
    
    try:
        json_script = data.get('json_script')
        project_id = data.get('project_id')
        
        if not json_script:
            raise HTTPException(status_code=400, detail="json_script is required")
        
        from src.services.voice_service import generate_voice_combined
        result = await generate_voice_combined(
            json_script=json_script,
            project_id=project_id,
        )
        
        if result.get('success'):
            print(f"✅ Combined voice generation complete: {result.get('total_slides')} slides → 1 audio file")
        else:
            print(f"⚠️ Combined voice generation failed: {result.get('error')}")
        
        return result
        
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in generate_voice_combined: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/enhance_prompts")
async def enhance_prompts_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Enhance visual cues from a script into detailed image generation prompts.
    
    Args:
        json_script: The parsed script JSON with slides containing image_prompt fields
        project_id: Optional project ID for tracking
    
    Returns:
        enhanced_prompts: List of {slide_number, title, original, enhanced, skip_reason}
    """
    print("🎨 Enhancing visual prompts...")
    
    try:
        json_script = data.get('json_script')
        project_id = data.get('project_id')
        
        if not json_script:
            raise HTTPException(status_code=400, detail="json_script is required")
        
        from src.services.prompt_enhancer import enhance_prompts
        result = enhance_prompts(json_script, project_id)
        
        print(f"✅ Enhanced {result.get('enhanced_count', 0)} prompts, skipped {result.get('skipped_count', 0)}")
        
        return result
        
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in enhance_prompts: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload_reference_image")
async def upload_reference_image(
    file: UploadFile = File(...),
    project_id: int = Form(...),
    slide_number: int = Form(...),
    current_user: TokenData = Depends(get_current_user)
):
    """
    Upload a reference image for image-to-image generation.
    
    Returns the path where the image was saved.
    """
    print(f"📎 Uploading reference image for slide {slide_number}...")
    
    try:
        # Create reference images directory
        project_root = Path(__file__).parent.parent.parent.parent
        ref_dir = project_root / "output" / "images" / str(project_id) / "references"
        ref_dir.mkdir(parents=True, exist_ok=True)
        
        # Save the file
        file_ext = Path(file.filename).suffix or ".png"
        ref_path = ref_dir / f"slide_{slide_number}_ref{file_ext}"
        
        content = await file.read()
        with open(ref_path, "wb") as f:
            f.write(content)
        
        print(f"  ✓ Saved reference image: {ref_path}")
        
        return {"path": str(ref_path), "slide_number": slide_number}
        
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in upload_reference_image: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate_images")
async def generate_images_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Generate images from approved prompts.
    
    Args:
        project_id: Project ID for file naming
        prompts: List of {slide_number, prompt} objects
        aspect_ratio: Optional, defaults to "1:1"
    
    Returns:
        images: List of {slide_number, url, success}
        zip_url: URL to download all images as ZIP
    """
    print("🖼️ Generating images from prompts...")
    
    try:
        project_id = data.get('project_id')
        prompts = data.get('prompts', [])
        aspect_ratio = data.get('aspect_ratio', '1:1')
        
        if not project_id:
            raise HTTPException(status_code=400, detail="project_id is required")
        
        if not prompts:
            raise HTTPException(status_code=400, detail="prompts list is required")
        
        from src.services.image_service import generate_images
        result = generate_images(prompts, project_id, aspect_ratio)
        
        print(f"✅ Generated {result.get('generated', 0)} images, {result.get('failed', 0)} failed")
        
        return result
        
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in generate_images: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate_slides")
async def generate_slides_endpoint(data: dict, current_user: TokenData = Depends(get_current_user)):
    """
    Generate a Beamer LaTeX template for presentation slides.
    
    Uses LLM-powered content extraction to intelligently parse scripts
    with context-aware intro phrasing (e.g., "learn to" vs "learn how to").
    
    Args:
        json_script: Optional - parsed script to extract content from
        tutorial_name: Name of the tutorial (optional, defaults to "Tutorial Name")
        num_content_slides: Number of blank content slides (optional, defaults to 10)
    
    Returns:
        tex_content: The complete LaTeX file content
        filename: Suggested filename for download
        zip_url: URL to download ZIP with .tex and assets
    """
    print("🎴 Generating Beamer slides template...")
    
    try:
        json_script = data.get('json_script')
        tutorial_name = data.get('tutorial_name', 'Tutorial Name')
        
        # Initialize template parameters
        template_params = {
            "tutorial_name": tutorial_name,
        }
        
        # Extract content from script using LLM if provided
        if json_script:
            print("🤖 Using LLM for intelligent content extraction...")
            
            from src.services.content_extractor import extract_slide_content_with_fallback
            extracted = extract_slide_content_with_fallback(json_script)
            
            # Update template params with extracted content
            template_params.update({
                "tutorial_name": extracted.get("tutorial_name", tutorial_name),
                "learning_objectives": extracted.get("learning_objectives"),
                "learning_objectives_intro": extracted.get("learning_objectives_intro"),
                "prerequisites": extracted.get("prerequisites"),
                "prerequisites_intro": extracted.get("prerequisites_intro"),
                "prerequisites_footer": extracted.get("prerequisites_footer"),
                "system_requirements": extracted.get("system_requirements"),
                "system_requirements_intro": extracted.get("system_requirements_intro"),
                "summary_points": extracted.get("summary_points"),
                "summary_intro": extracted.get("summary_intro"),
                "assignment_items": extracted.get("assignment_items"),
                "assignment_intro": extracted.get("assignment_intro"),
                "domain_expert": extracted.get("domain_expert"),
                "domain_expert_org": extracted.get("domain_expert_org"),
                "code_file_info": extracted.get("code_file_info"),
            })
            
            print(f"📝 LLM Extracted: LO={extracted.get('learning_objectives') is not None}, "
                  f"Prereq={extracted.get('prerequisites') is not None}, "
                  f"Summary={extracted.get('summary_points') is not None}, "
                  f"Assignment={extracted.get('assignment_items') is not None}")
            print(f"📝 Intro phrases: LO='{extracted.get('learning_objectives_intro', '')[:50]}...'")
        
        from src.services.beamer_service import generate_beamer_template
        tex_content = generate_beamer_template(**template_params)
        
        # Generate a safe filename
        safe_name = template_params["tutorial_name"].replace(' ', '_').replace('/', '-')[:50]
        tex_filename = f"{safe_name}_slides.tex"
        zip_filename = f"{safe_name}_slides.zip"
        
        # Create ZIP file with .tex and logo
        import zipfile
        
        # Get project root for static assets
        project_root = Path(__file__).parent.parent.parent.parent
        logo_path = project_root / "static" / "logo.png"
        
        # Create output directory if needed
        output_dir = project_root / "output" / "slides"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        zip_path = output_dir / zip_filename
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # Add the .tex file
            zipf.writestr(tex_filename, tex_content)
            
            # Add the logo if it exists
            if logo_path.exists():
                zipf.write(logo_path, "logo.png")
            else:
                print(f"⚠️ Logo not found at {logo_path}")
        
        auto_filled = json_script is not None
        print(f"✅ Generated Beamer ZIP: {zip_filename}" + (" (LLM-extracted)" if auto_filled else ""))
        
        return {
            "tex_content": tex_content,
            "filename": tex_filename,
            "zip_filename": zip_filename,
            "zip_url": f"/output/slides/{zip_filename}",
            "num_boilerplate_slides": 8,  # Title, LO, SysReq, Prereq, Code, Summary, Assignment, Thanks
            "auto_filled": auto_filled
        }
        
    except Exception as e:
        traceback.print_exc()
        print(f"ERROR in generate_slides: {e}")
        raise HTTPException(status_code=500, detail=str(e))




