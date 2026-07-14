from io import BytesIO
import unittest

from docx import Document

from src.services.docx_service import docx_to_json


class DocxServiceTests(unittest.TestCase):
    def test_docx_to_json_extracts_hyphenated_prerequisite_metadata_without_script_overwrite(self):
        doc = Document()

        metadata = doc.add_table(rows=0, cols=2)
        rows = [
            ("Series", "TensorFlow"),
            ("Tutorial:", "4. Tensor Operations"),
            ("Approx. Duration:", "4-5 mins"),
            ("Pre-requisite Tutorial", "Basic Python syntax and TensorFlow constants."),
            ("Meta Tags", "Tensor Operations, TensorFlow"),
            ("Outline", "Add tensors\nSubtract tensors"),
        ]
        for label, value in rows:
            row = metadata.add_row()
            row.cells[0].text = label
            row.cells[1].text = value

        script = doc.add_table(rows=3, cols=2)
        script.rows[0].cells[0].text = "Visual Cue"
        script.rows[0].cells[1].text = "Narration"
        script.rows[1].cells[0].text = "Slide 1\nTitle Slide"
        script.rows[1].cells[1].text = "Welcome to the Spoken Tutorial on Tensor Operations."
        script.rows[2].cells[0].text = "Slide 5\nPre-requisite slide\nhttp://EduPyramids.org"
        script.rows[2].cells[1].text = "For the pre-requisite of this tutorial, please visit the website shown on your screen."

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        parsed = docx_to_json(buffer)

        self.assertEqual(parsed["presentation_title"], "Tensor Operations")
        self.assertEqual(parsed["tutorial"], "4. Tensor Operations")
        self.assertEqual(parsed["prerequisites"], "Basic Python syntax and TensorFlow constants.")
